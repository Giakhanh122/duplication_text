import gradio as gr
import streamlit as st
import pandas as pd
import docx

import re
from ftfy import fix_text
from FaissSearch import *
from HSmodule import *
from sentence_transformers import SentenceTransformer


def normalizing(text: str):
    text = fix_text(text)
    text = re.sub(r"\s+", " ", text).strip()
    text = text.lower()
    return text


def read_file(filepath):
    #tach file extension
    ext = os.path.splitext(filepath)[1].lower()
    # tach du lieu tu file
    if ext == ".docx":
      doc = docx.Document(filepath)
      paragraphs = [normalizing(p.text.strip()) for p in doc.paragraphs if p.text.strip()]
    elif ext == ".txt":
      with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()
      paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
    elif ext == ".csv":
      table = pd.read_csv(filepath)
      data_cols = ["content", "text", "paragraph"]
      for col in data_cols:
        if col in table.columns:
           paragraphs = table[col].dropna().astype(str).tolist()
           break
    else:
        raise ValueError("Không tìm thấy cột 'text' / 'content' / 'paragraph' trong file CSV hoặc file không hợp lệ.")
    return paragraphs



def run_SimHash(paragraphs : list[str]) -> list[list[VectorRecord]]:
    model = SentenceTransformer('sentence-transformers/all-MiniLM-L6-v2')
    # sinh vector
    embedding = model.encode(paragraphs)
    # tao list vectorRecord
    list_vectors = [VectorRecord(id, vec) for id, vec in enumerate(embedding)]
    # setup Simhash
    hasher = SimHash(InputDim = 384 , OutputDim = 128)
    # setup Search
    searcher =  LSHSearch()
    searcher.bandSize = 8
    searcher.threshold = 0.25
    searcher.setDisFunc("hamming")
    # hash
    hashed_vectors = hasher.hash(list_vectors)
    # classify
    ans = searcher.classify(hashed_vectors)
    return ans


def run_Faiss(paragraphs : list[str]) -> list[list[VectorRecord]]:
  model = SentenceTransformer('sentence-transformers/all-MiniLM-L6-v2')
  # sinh vector
  embedding = model.encode(paragraphs)
  # tao list vectorRecord
  list_vectors = [VectorRecord(id, vec) for id, vec in enumerate(embedding)]
  # setup Simhash
  hasher = BloomFilter(0, 0, expectedItems=1000, falsePositiveRate=0.01)
  # setup Search
  searcher =  FaissSearch()
  searcher.threshold = 0.2
  searcher.setDisFunc("cosine")
  # hash
  hashed_vectors = hasher.hash(list_vectors)
  # classify
  ans = searcher.classify(hashed_vectors)
  return ans


# triển khai thêm 1 hướng ??

def duplication_text(filepath : str, method : str):

    # tach du lieu
    paragraphs = read_file(filepath)
    if method == "SimHash":
      ans = run_SimHash(paragraphs)
    elif method == "Bloom + Faiss":
      ans = run_Faiss(paragraphs)
    #hightlight
    # màu để highlight - 17 màu
    from docx.enum.text import WD_COLOR_INDEX
    colors = [
        None,  # 0 → không highlight
        WD_COLOR_INDEX.BLACK,
        WD_COLOR_INDEX.BLUE,
        WD_COLOR_INDEX.BRIGHT_GREEN,
        WD_COLOR_INDEX.DARK_BLUE,
        WD_COLOR_INDEX.DARK_RED,
        WD_COLOR_INDEX.DARK_YELLOW,
        WD_COLOR_INDEX.GRAY_25,
        WD_COLOR_INDEX.GRAY_50,
        WD_COLOR_INDEX.GREEN,
        WD_COLOR_INDEX.PINK,
        WD_COLOR_INDEX.RED,
        WD_COLOR_INDEX.TEAL,
        WD_COLOR_INDEX.TURQUOISE,
        WD_COLOR_INDEX.VIOLET,
        WD_COLOR_INDEX.WHITE,
        WD_COLOR_INDEX.YELLOW,
    ]


    # tạo dictionraty để truy xuất xử lý dữ liệu
    group_id = {x.id: (0 if len(group) == 1 else group_index) for group_index, group in enumerate(ans) for x in group}

    # tạo document mới

    # doc1 = docx.Document()
    # doc1.add_heading("Duplicate Text Grouping Result:", level=1)
    # for id, group in enumerate(ans):
    #   doc1.add_heading(f"Group {id} : ", level=2)
    #   for para in group:
    #     index = para.id
    #     doc1.add_paragraph(f"Paragraph {index} : {paragraphs[index]}")
    # grouped_file = "grouped.docx"
    # doc1.save(grouped_file)


    doc2 = docx.Document()
    doc2.add_heading("Duplicate Text Highlighting Result:", level=1)
    for id, para in enumerate(paragraphs):
        if id in group_id:
          p = doc2.add_paragraph()
          run = p.add_run("    ")
          color = colors[group_id[id] % len(colors)]
          run.font.highlight_color = color
          p.add_run(f"Paragraph {id} : {para}")
    highlighted_file = "highlighted.docx"
    doc2.save(highlighted_file)
    # return highlighted_file, grouped_file

    # in ket qua
    html = "<h2>Duplicate Text Grouping Result</h2>"
    for id, group in enumerate(ans):
      html += f"<h3 style='color:#0af'>Group {id} ({len(group)} items)</h3><ul>"
      for para in group:
        html += f"<li><b>Paragraph {para.id}</b>: {paragraphs[para.id][:300]}...</li>"
      html += "</ul>"

    return html, highlighted_file







# dùng gradio để tạo giao diện demo
# with gr.Blocks(title="Duplicate Text Detector") as demo:
#     gr.Markdown("## 🧩 Duplicate Text Detector\nUpload a `.docx`, `.txt`, or `.csv` file to detect similar paragraphs using SimHash or Bloom+FAISS.")

#     with gr.Row():  # vùng chứa input
#         with gr.Column():
#             file_input = gr.File(label="Upload document file (.docx, .txt, .csv)")
#             method_choice = gr.Radio(
#                 ["SimHash", "Bloom + Faiss"],
#                 label="Choose method :",
#                 value="SimHash"
#             )
#             submit_btn = gr.Button("Submit", variant="primary")

#     with gr.Column():  # vùng chứa output (nằm DƯỚI)
#         html_output = gr.HTML(label="Duplicate text result")
#         file_output = gr.File(label="Download grouped result")

#     # kết nối logic
#     submit_btn.click(
#         fn=duplication_text,
#         inputs=[file_input, method_choice],
#         outputs=[html_output, file_output]
#     )

# demo.launch()


#streamlit
# UI Streamlit
uploaded_file = st.file_uploader("Upload document file (.docx, .txt, .csv)")
method = st.radio("Choose method:", ["SimHash", "Bloom + Faiss"])
if st.button("Submit") and uploaded_file is not None:
    with st.spinner("Processing..."):
        # lưu tạm file
        filepath = f"temp_{uploaded_file.name}"
        with open(filepath, "wb") as f:
            f.write(uploaded_file.getbuffer())

        html_result, highlighted_file = duplication_text(filepath, method)
        
        # hiển thị kết quả HTML
        st.markdown(html_result, unsafe_allow_html=True)
        # link download
        st.download_button("Download highlighted docx", highlighted_file)

